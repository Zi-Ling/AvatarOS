# app/avatar/skills/builtin/word.py

from __future__ import annotations

import os
from pydantic import Field, model_validator
from docx import Document
from pathlib import Path
from ..common.path_normalizer import normalize_file_extension
from ..common.path_mixins import PathBindMixin
from ..common.content_mixins import ContentRobustnessMixin

from ..base import BaseSkill, SkillSpec, SkillCategory, SkillPermission, SkillMetadata, SkillDomain, SkillCapability
from ..schema import SkillInput, SkillOutput
from ..registry import register_skill
from ..context import SkillContext


def _load_or_create_document(path):
    try:
        return Document(path)
    except Exception:
        return Document()


# ============================================================================
# word.write_text
# ============================================================================

class WordWriteTextInput(PathBindMixin, ContentRobustnessMixin, SkillInput):
    # relative_path 可以允许为空，因为有可能完全靠 file_path/abs_path 驱动
    relative_path: str | None = Field(
        None, description="Relative Word file path (should end with .docx)."
    )
    content: str = Field(..., description="Full text content to write into the document.")
    
    # 可选：增加一个 abs_path，直接使用绝对路径时用
    abs_path: str | None = Field(
        None, description="Absolute file path. If provided, takes precedence."
    )

    @model_validator(mode="after")
    def normalize_ext(self):
        if self.relative_path:
            self.relative_path = normalize_file_extension(
                self.relative_path, 
                default_ext=".docx", 
                allowed_exts={".docx", ".doc"}
            )
        return self

class WordWriteTextOutput(SkillOutput):
    path: str
    content_length: int

@register_skill
class WordWriteTextSkill(BaseSkill[WordWriteTextInput, WordWriteTextOutput]):
    spec = SkillSpec(
        name="word.write_text",
        api_name="word.write",
        aliases=["docx.write", "write_docx"],
        description="Create or overwrite a .docx file with text content. 创建或覆盖Word文档(.docx)。",
        category=SkillCategory.OFFICE,
        input_model=WordWriteTextInput,
        output_model=WordWriteTextOutput,
        
        # Capability Routing
        meta=SkillMetadata(
            domain=SkillDomain.OFFICE,
            capabilities={SkillCapability.WRITE, SkillCapability.CREATE},
            risk_level="high",
            file_extensions=[".docx", ".doc"]
        ),
        
        # ✅ Artifact Management（混合方案：声明式）
        produces_artifact=True,
        artifact_type="document:word",
        artifact_path_field="path",
        
        permissions=[SkillPermission(name="file_write", description="Write Word files")],
        tags=["office", "word", "docx", "文档", "Word", "保存"]
    )

    async def run(self, ctx: SkillContext, params: WordWriteTextInput) -> WordWriteTextOutput:
        # 1. 优先使用 abs_path
        if params.abs_path:
             target_path = Path(params.abs_path)
        # 2. 否则使用 relative_path
        elif params.relative_path:
             target_path = ctx.resolve_path(params.relative_path)
        # 3. 如果都没有，报错
        else:
             return WordWriteTextOutput(success=False, message="No valid path provided (neither relative_path nor abs_path).", path="", content_length=0)

        if ctx.dry_run:
            return WordWriteTextOutput(
                success=True,
                message=f"[dry_run] Write Docx: {target_path}",
                path=str(target_path),
                content_length=len(params.content)
            )

        try:
            target_path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            for para in params.content.split("\n"):
                doc.add_paragraph(para)
            doc.save(str(target_path))
            return WordWriteTextOutput(
                success=True,
                message=f"Written Docx: {target_path}",
                path=str(target_path),
                content_length=len(params.content),
                # FS Metadata
                fs_operation='created',
                fs_path=params.relative_path or os.path.basename(params.abs_path),
                fs_type='file'
            )
        except Exception as e:
            return WordWriteTextOutput(success=False, message=str(e), path=str(target_path), content_length=0)


# ============================================================================
# word.append_text
# ============================================================================

class WordAppendTextInput(PathBindMixin, ContentRobustnessMixin, SkillInput):
    # relative_path 可以允许为空，因为有可能完全靠 file_path/abs_path 驱动
    relative_path: str | None = Field(
        None, description="Relative Word file path."
    )
    content: str = Field(..., description="Text to append as new paragraphs.")

    # 可选：增加一个 abs_path，直接使用绝对路径时用
    abs_path: str | None = Field(
        None, description="Absolute file path. If provided, takes precedence."
    )

    @model_validator(mode="after")
    def normalize_ext(self):
        if self.relative_path:
            self.relative_path = normalize_file_extension(
                self.relative_path, 
                default_ext=".docx", 
                allowed_exts={".docx", ".doc"}
            )
        return self

class WordAppendTextOutput(SkillOutput):
    path: str
    content_length: int

@register_skill
class WordAppendTextSkill(BaseSkill[WordAppendTextInput, WordAppendTextOutput]):
    spec = SkillSpec(
        name="word.append_text",
        api_name="word.append",
        aliases=["docx.append", "append_docx"],
        description="Append text to an existing Word document. 在现有Word文档中追加文本内容。",
        category=SkillCategory.OFFICE,
        input_model=WordAppendTextInput,
        output_model=WordAppendTextOutput,
        
        # Capability Routing
        meta=SkillMetadata(
            domain=SkillDomain.OFFICE,
            capabilities={SkillCapability.WRITE, SkillCapability.MODIFY},
            risk_level="normal",
            file_extensions=[".docx", ".doc"]
        ),
        
        # ✅ Artifact Management（混合方案：声明式）
        produces_artifact=True,
        artifact_type="document:word",
        artifact_path_field="path",
        
        permissions=[SkillPermission(name="file_write", description="Modify Word files")],
        tags=["office", "word", "append", "文档", "追加", "Word"]
    )

    async def run(self, ctx: SkillContext, params: WordAppendTextInput) -> WordAppendTextOutput:
        # 1. 优先使用 abs_path
        if params.abs_path:
             target_path = Path(params.abs_path)
        # 2. 否则使用 relative_path
        elif params.relative_path:
             target_path = ctx.resolve_path(params.relative_path)
        # 3. 如果都没有，报错
        else:
             return WordAppendTextOutput(success=False, message="No valid path provided (neither relative_path nor abs_path).", path="", content_length=0)

        if ctx.dry_run:
            return WordAppendTextOutput(
                success=True,
                message=f"[dry_run] Append to Docx: {target_path}",
                path=str(target_path),
                content_length=len(params.content)
            )

        try:
            target_path.parent.mkdir(parents=True, exist_ok=True)
            doc = _load_or_create_document(str(target_path))
            for para in params.content.split("\n"):
                doc.add_paragraph(para)
            doc.save(str(target_path))
            return WordAppendTextOutput(
                success=True,
                message=f"Appended to Docx: {target_path}",
                path=str(target_path),
                content_length=len(params.content),
                # FS Metadata
                fs_operation='modified',
                fs_path=params.relative_path or os.path.basename(params.abs_path),
                fs_type='file'
            )
        except Exception as e:
            return WordAppendTextOutput(success=False, message=str(e), path=str(target_path), content_length=0)
